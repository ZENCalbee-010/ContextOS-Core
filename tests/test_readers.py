from pathlib import Path

import fitz
from docx import Document

from contextos.readers.code_reader import CodeReader
from contextos.readers.detector import get_reader, read_file
from contextos.readers.docx_reader import DOCXReader
from contextos.readers.pdf_reader import PDFReader
from contextos.readers.text_reader import TextReader


def test_detector_selects_reader_by_extension(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    docx_path = tmp_path / "sample.docx"
    text_path = tmp_path / "sample.txt"
    markdown_path = tmp_path / "sample.md"
    code_path = tmp_path / "sample.py"

    for path in [pdf_path, docx_path, text_path, markdown_path, code_path]:
        path.touch()

    assert isinstance(get_reader(pdf_path), PDFReader)
    assert isinstance(get_reader(docx_path), DOCXReader)
    assert isinstance(get_reader(text_path), TextReader)
    assert isinstance(get_reader(markdown_path), TextReader)
    assert isinstance(get_reader(code_path), CodeReader)


def test_text_reader_returns_text_and_metadata(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("# ContextOS\n\nSelection over compression.", encoding="utf-8")

    result = read_file(path)

    assert result.supported is True
    assert result.error is None
    assert "Selection over compression" in result.text
    assert result.metadata["extension"] == ".md"
    assert result.metadata["format"] == "markdown"
    assert result.metadata["reader"] == "text"


def test_code_reader_returns_text_and_language_metadata(tmp_path):
    path = tmp_path / "app.ts"
    path.write_text("const answer: number = 42;\n", encoding="utf-8")

    result = read_file(path)

    assert result.supported is True
    assert "const answer" in result.text
    assert result.metadata["language"] == "typescript"
    assert result.metadata["reader"] == "code"


def test_pdf_reader_returns_text_and_metadata(tmp_path):
    path = tmp_path / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF context sample")
    document.save(path)
    document.close()

    result = read_file(path)

    assert result.supported is True
    assert "PDF context sample" in result.text
    assert result.metadata["page_count"] == 1
    assert result.metadata["reader"] == "pdf"


def test_docx_reader_returns_text_and_metadata(tmp_path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("DOCX context sample")
    document.add_paragraph("Second paragraph")
    document.save(path)

    result = read_file(path)

    assert result.supported is True
    assert "DOCX context sample" in result.text
    assert "Second paragraph" in result.text
    assert result.metadata["paragraph_count"] == 2
    assert result.metadata["reader"] == "docx"


def test_unsupported_file_is_handled_gracefully(tmp_path):
    path = tmp_path / "archive.zip"
    path.write_bytes(b"not a real zip")

    result = read_file(path)

    assert result.supported is False
    assert result.text == ""
    assert result.error == "Unsupported file type: .zip"
    assert result.metadata["extension"] == ".zip"


def test_missing_supported_file_returns_error_result(tmp_path):
    path = tmp_path / "missing.txt"

    result = read_file(path)

    assert result.supported is True
    assert result.text == ""
    assert result.error is not None


def test_all_supported_code_extensions_have_readers(tmp_path):
    for extension in [".py", ".js", ".ts", ".html", ".css", ".json"]:
        path = Path(tmp_path / f"sample{extension}")
        path.write_text("sample", encoding="utf-8")

        result = read_file(path)

        assert result.supported is True
        assert result.metadata["reader"] == "code"
