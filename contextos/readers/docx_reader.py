# ---------------------------------------------------------------------------
# readers/docx_reader.py — Reader สำหรับอ่านไฟล์ Microsoft Word (.docx)
# ---------------------------------------------------------------------------
# หน้าที่ : อ่านเนื้อหาจากไฟล์ .docx แล้วแปลงเป็นข้อความธรรมดา (plain text)
# ความรับผิดชอบ :
#   - ดึงข้อความจากทุก paragraph ที่ไม่ว่าง
#   - ดึง metadata จาก core_properties ของเอกสาร (title, author)
# ความสัมพันธ์กับระบบ :
#   - เป็น Concrete Strategy ของ BaseReader (Strategy Pattern)
#   - ผลลัพธ์จะถูกส่งต่อให้ TextParser เพื่อตัดเป็น chunks
# หมายเหตุ : ใช้ไลบรารี python-docx ซึ่งรองรับเฉพาะ .docx (Open XML)
#   ไม่รองรับ .doc (Binary Format ของ Word เก่า)
# ---------------------------------------------------------------------------
"""DOCX reader powered by python-docx."""

from pathlib import Path

# python-docx — ไลบรารีสำหรับอ่าน/เขียนไฟล์ .docx
# ใช้ชื่อ DocxDocument เพื่อไม่ให้ชนกับ Document ตัวอื่นในโปรเจกต์
from docx import Document as DocxDocument

from contextos.readers.base import BaseReader, ReaderResult


# ---------------------------------------------------------------------------
# DOCXReader — Concrete Strategy สำหรับอ่านไฟล์ Word
# ---------------------------------------------------------------------------
# ใช้ทำอะไร : แปลงไฟล์ .docx เป็นข้อความธรรมดาพร้อม metadata
# รับผิดชอบอะไร : ดึงข้อความจาก paragraphs + ข้อมูลเอกสาร
# ใช้เมื่อไร : เมื่อ detector เลือก reader นี้จากนามสกุล .docx
# ทำงานร่วมกับ : BaseReader (parent), detector.py (เลือก reader)
class DOCXReader(BaseReader):
    supported_extensions = frozenset({".docx"})
    reader_name = "docx"

    # จุดประสงค์ : อ่านไฟล์ .docx และสร้าง ReaderResult
    # Input : path — เส้นทางไฟล์ .docx ที่จะอ่าน
    # Output : ReaderResult พร้อมข้อความจาก paragraph และ metadata
    # ขั้นตอนการทำงาน :
    #   1. เปิดไฟล์ด้วย python-docx
    #   2. ดึงข้อความจาก paragraph ที่ไม่ว่าง
    #   3. ดึง metadata พื้นฐาน + เพิ่มจำนวน paragraph, title, author
    def read(self, path: str | Path) -> ReaderResult:
        file_path = Path(path)
        document = DocxDocument(file_path)
        # กรอง paragraph ที่ว่างออก เพราะ Word มักมี paragraph เปล่าระหว่างเนื้อหา
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

        metadata = self._base_metadata(file_path)
        # เพิ่ม metadata เฉพาะทางของ DOCX : จำนวน paragraph, ชื่อเรื่อง, ผู้เขียน
        # ใช้ `or None` เพื่อแปลง string ว่างเป็น None ให้สอดคล้องกับ convention
        metadata.update(
            {
                "paragraph_count": len(document.paragraphs),
                "title": document.core_properties.title or None,
                "author": document.core_properties.author or None,
            }
        )

        # รวม paragraphs ด้วย newline เพื่อรักษาโครงสร้างเอกสาร
        return ReaderResult(text="\n".join(paragraphs), metadata=metadata)
